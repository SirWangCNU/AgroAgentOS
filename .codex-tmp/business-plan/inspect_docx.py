from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


DOCX = Path(r"E:\GithubProgram\AgroAgentOS\.codex-tmp\business-plan\template.docx")


def iter_block_items(parent):
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


doc = Document(DOCX)
print(f"PARAGRAPHS={len(doc.paragraphs)} TABLES={len(doc.tables)} SECTIONS={len(doc.sections)}")
for index, block in enumerate(iter_block_items(doc), 1):
    if isinstance(block, Paragraph):
        text = block.text.replace("\t", "<TAB>").strip()
        if text:
            print(f"P{index:03d} [{block.style.name}] {text}")
    else:
        print(f"TABLE{index:03d} rows={len(block.rows)} cols={len(block.columns)}")
        for row_index, row in enumerate(block.rows, 1):
            cells = [" / ".join(p.text.strip() for p in cell.paragraphs if p.text.strip()) for cell in row.cells]
            print(f"  R{row_index:02d}: " + " || ".join(cells))

for section_index, section in enumerate(doc.sections, 1):
    for kind, container in (("HEADER", section.header), ("FOOTER", section.footer)):
        texts = [p.text.strip() for p in container.paragraphs if p.text.strip()]
        if texts:
            print(f"{kind}{section_index}: " + " | ".join(texts))
