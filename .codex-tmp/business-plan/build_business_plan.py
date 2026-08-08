from __future__ import annotations

import copy
import hashlib
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\GithubProgram\AgroAgentOS")
TMP = ROOT / ".codex-tmp" / "business-plan"
REFERENCE = TMP / "template.docx"
CONTENT = TMP / "business_plan_content.md"
FIGURES = TMP / "figures"
OUTPUT = ROOT / "output" / "AgroAgentOS_第十五届挑战杯商业计划书_更新版.docx"
EXPECTED_HASH = "AC19A0CE653793FCF7E3D99FF8C511636C8257B3B9372568EFC8609A5BD406C5"

BODY_FONT = "宋体"
WESTERN_FONT = "Times New Roman"
H1_FONT = "黑体"
H2_FONT = "楷体"
ACCENT = "0B3558"
TABLE_FILL = "DCEAF0"
USABLE_DXA = 8720


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


def set_run_font(run, east_asia=BODY_FONT, western=WESTERN_FONT, size=14, bold=None, color=None):
    run.font.name = western
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), western)


def set_style_font(style, east_asia, western, size, bold=False, color=None):
    style.font.name = western
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    set_style_font(normal, BODY_FONT, WESTERN_FONT, 14, False)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(28)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(28)
    pf.widow_control = True

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, H1_FONT, WESTERN_FONT, 16, False, ACCENT)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.first_line_indent = Pt(0)
    h1.paragraph_format.left_indent = Pt(0)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h1.paragraph_format.line_spacing = Pt(28)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, H2_FONT, WESTERN_FONT, 16, False, ACCENT)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.first_line_indent = Pt(0)
    h2.paragraph_format.left_indent = Pt(0)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h2.paragraph_format.line_spacing = Pt(28)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, BODY_FONT, WESTERN_FONT, 14, True, ACCENT)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.first_line_indent = Pt(0)
    h3.paragraph_format.left_indent = Pt(0)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h3.paragraph_format.line_spacing = Pt(28)
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(0)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.keep_together = True

    for list_name in ("List Bullet", "List Number"):
        try:
            style = doc.styles[list_name]
        except KeyError:
            continue
        set_style_font(style, BODY_FONT, WESTERN_FONT, 14, False)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(28)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.first_line_indent = Pt(0)


def clear_keep_ppr(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        clear_keep_ppr(paragraph)
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._p
        p.getparent().remove(p)
    return cell.paragraphs[0]


def update_cover(doc: Document):
    p = doc.paragraphs[9]
    clear_keep_ppr(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    label = p.add_run("项目名称：  ")
    set_run_font(label, "微软雅黑", WESTERN_FONT, 14, True, "FFFFFF")
    name = p.add_run("AgroAgentOS——面向中小农场的多智能体智慧农服平台")
    set_run_font(name, "微软雅黑", WESTERN_FONT, 11.5, True, "FFFFFF")

    table = doc.tables[0]
    p1 = clear_cell(table.cell(0, 0))
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("赛      道：")
    set_run_font(r1, "微软雅黑", WESTERN_FONT, 12, True, "FFFFFF")
    p2 = clear_cell(table.cell(0, 1))
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("乡村振兴和农业农村现代化")
    set_run_font(r2, "微软雅黑", WESTERN_FONT, 11.5, True, "FFFFFF")


def add_toc_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录将在Microsoft Word中自动更新"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, BODY_FONT, WESTERN_FONT, 14, False)


def add_section_break_paragraph(doc: Document, sect_pr):
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    ppr.append(copy.deepcopy(sect_pr))
    return p


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def visual_len(text: str) -> float:
    return sum(1.0 if ord(ch) > 127 else 0.55 for ch in text)


def table_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    weights = []
    for i in range(cols):
        maximum = max(visual_len(row[i]) for row in rows)
        weights.append(max(5.5, min(maximum, 24.0)))
    total = sum(weights)
    widths = [int(USABLE_DXA * weight / total) for weight in weights]
    minimum = 820 if cols >= 5 else 980
    for i, value in enumerate(widths):
        widths[i] = max(value, minimum)
    while sum(widths) > USABLE_DXA:
        index = max(range(cols), key=lambda j: widths[j])
        if widths[index] <= minimum:
            break
        widths[index] -= min(40, sum(widths) - USABLE_DXA)
    widths[-1] += USABLE_DXA - sum(widths)
    return widths


def apply_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    layout = tblPr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[i]))
            tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[i] / 1440)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trPr.append(header)


def ensure_bullet_numbering(doc: Document) -> int:
    cached = getattr(doc, "_bp_bullet_num_id", None)
    if cached is not None:
        return cached
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "560")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "560")
    ind.set(qn("w:hanging"), "280")
    ppr.extend([tabs, ind])
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Symbol")
    fonts.set(qn("w:hAnsi"), "Symbol")
    rpr.append(fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, ppr, rpr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    setattr(doc, "_bp_bullet_num_id", num_id)
    return num_id


def apply_bullet(paragraph, num_id: int):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    ppr.append(num_pr)


def add_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = table_widths(rows)
    font_size = 10.5 if len(rows[0]) >= 5 else 11.5 if len(rows[0]) == 4 else 12
    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = clear_cell(cell)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(font_size + 5)
            p.paragraph_format.keep_together = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if visual_len(text) > 13 and r_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, BODY_FONT, WESTERN_FONT, font_size, r_idx == 0, ACCENT if r_idx == 0 else None)
            if r_idx == 0:
                shade_cell(cell, TABLE_FILL)
    repeat_header(table.rows[0])
    apply_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    after.paragraph_format.line_spacing = Pt(8)
    return table


def add_table_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_run_font(r, H1_FONT, WESTERN_FONT, 12, False, ACCENT)


def add_figure(doc: Document, filename: str, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    # The body style uses an exact 28 pt line height. Inline pictures inherit
    # that setting unless the picture paragraph explicitly opts into an
    # auto-expanding line height, which would crop every figure in Word.
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(FIGURES / filename), width=Inches(5.85))
    shape._inline.docPr.set("title", caption)
    shape._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Pt(0)
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_together = True
    r = cap.add_run(caption)
    set_run_font(r, H1_FONT, WESTERN_FONT, 12, False, ACCENT)


def add_body_paragraph(doc: Document, text: str):
    is_bullet = False
    if text.startswith("- "):
        is_bullet = True
        text = text[2:].strip()
    p = doc.add_paragraph(style="Normal")
    if is_bullet:
        apply_bullet(p, ensure_bullet_numbering(doc))
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(28)
        p.paragraph_format.keep_together = True
    if re.match(r"^\[\d+\]", text):
        p.paragraph_format.first_line_indent = Pt(-20)
        p.paragraph_format.left_indent = Pt(20)
    r = p.add_run(text)
    set_run_font(r, BODY_FONT, WESTERN_FONT, 14, False)
    return p


def parse_content(doc: Document, markdown: str):
    lines = markdown.splitlines()
    i = 0
    h1_count = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith("[[FIG:"):
            match = re.match(r"\[\[FIG:([^|]+)\|(.+)\]\]", line)
            if not match:
                raise ValueError(f"Bad figure marker: {line}")
            add_figure(doc, match.group(1), match.group(2))
            i += 1
            continue
        if line.startswith("# "):
            h1_count += 1
            p = doc.add_paragraph(line[2:].strip(), style="Heading 1")
            p.paragraph_format.page_break_before = h1_count > 1
            i += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 2")
            i += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 3")
            i += 1
            continue
        if line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            parsed = [[cell.strip() for cell in row.strip("|").split("|")] for row in block]
            parsed = [row for idx, row in enumerate(parsed) if idx != 1 or not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in row)]
            add_table(doc, parsed)
            continue
        if re.match(r"^表\d+\.\d+", line):
            add_table_caption(doc, line)
            i += 1
            continue
        add_body_paragraph(doc, line)
        i += 1


def set_update_fields(doc: Document):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def normalize_footer_page_number_boxes(doc: Document):
    """Keep the template's floating page-number boxes usable past page 9."""
    seen_parts = set()
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            part_name = str(footer.part.partname)
            if part_name in seen_parts:
                continue
            seen_parts.add(part_name)
            root = footer._element
            for shape in root.xpath(".//*[local-name()='shape']"):
                style = shape.get("style", "")
                if "width:" in style:
                    style = re.sub(r"width:[^;]+", "width:2in", style, count=1)
                else:
                    style = f"{style};width:2in"
                shape.set("style", style)
            for paragraph in root.xpath(
                ".//*[local-name()='txbxContent']/*[local-name()='p']"
            ):
                ppr = paragraph.find(qn("w:pPr"))
                if ppr is None:
                    ppr = OxmlElement("w:pPr")
                    paragraph.insert(0, ppr)
                ind = ppr.find(qn("w:ind"))
                if ind is not None:
                    ind.set(qn("w:firstLine"), "0")
                    ind.set(qn("w:left"), "0")
                jc = ppr.find(qn("w:jc"))
                if jc is None:
                    jc = OxmlElement("w:jc")
                    ppr.append(jc)
                jc.set(qn("w:val"), "center")


def replace_body_footer_with_page_field(doc: Document):
    """Use a normal centered PAGE field so two-digit numbers cannot be clipped."""
    footer = doc.sections[-1].footer
    footer.is_linked_to_previous = False
    root = footer._element
    for child in list(root):
        root.remove(child)

    paragraph = footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE \\* MERGEFORMAT "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, BODY_FONT, WESTERN_FONT, 12, False)


def main():
    if sha256(REFERENCE) != EXPECTED_HASH:
        raise RuntimeError("Reference template hash mismatch; re-distill before building")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(REFERENCE)
    configure_styles(doc)
    update_cover(doc)

    body = doc.element.body
    children = list(body.iterchildren())
    toc_sect_pr = children[47].find(".//" + qn("w:sectPr"))
    body_sect_pr = children[138].find(".//" + qn("w:sectPr"))
    if toc_sect_pr is None or body_sect_pr is None:
        raise RuntimeError("Template section locators changed")
    for child in children[23:]:
        body.remove(child)

    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.first_line_indent = Pt(0)
    toc_title.paragraph_format.space_before = Pt(6)
    toc_title.paragraph_format.space_after = Pt(18)
    toc_title.paragraph_format.keep_with_next = True
    r = toc_title.add_run("目    录")
    set_run_font(r, H1_FONT, WESTERN_FONT, 18, False, ACCENT)
    toc = doc.add_paragraph()
    toc.paragraph_format.first_line_indent = Pt(0)
    add_toc_field(toc)
    add_section_break_paragraph(doc, toc_sect_pr)

    parse_content(doc, CONTENT.read_text(encoding="utf-8"))
    body.append(copy.deepcopy(body_sect_pr))

    set_update_fields(doc)
    normalize_footer_page_number_boxes(doc)
    replace_body_footer_with_page_field(doc)
    props = doc.core_properties
    props.title = "AgroAgentOS——面向中小农场的多智能体智慧农服平台"
    props.subject = "第十五届“挑战杯”中国大学生创业计划竞赛商业计划书"
    props.author = "AgroAgentOS项目团队"
    props.keywords = "挑战杯,智慧农业,多智能体,农业数字化,商业计划书"
    props.comments = "依据第十五届挑战杯商业计划书模板生成；财务数据为预测口径。"
    doc.save(OUTPUT)
    print(OUTPUT)
    print(f"SIZE={OUTPUT.stat().st_size}")


if __name__ == "__main__":
    main()
